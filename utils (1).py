"""
Shared text-processing + matching utilities for the Resume Screening System.

Two cleaning paths are kept, matching the two things the model was trained
around in the original notebook:

1. clean_text()            -> simple clean/lemmatize, used to prep text
                               before TF-IDF for the SVM classifier.
2. clean_text_negation()   -> same idea but tags words that follow a
                               negation cue (e.g. "not fluent in Java" ->
                               "NEG_fluent NEG_in NEG_java"), used for the
                               JD-vs-resume matching/ranking feature so that
                               "no experience with X" doesn't score as a
                               match for X.
"""

import io
import os
import re

import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"}


def ensure_nltk_data() -> None:
    """Download required NLTK corpora if they aren't already present."""
    resources = {
        "corpora/stopwords": "stopwords",
        "corpora/wordnet": "wordnet",
        "tokenizers/punkt": "punkt",
        "tokenizers/punkt_tab": "punkt_tab",
    }
    for path, name in resources.items():
        try:
            nltk.data.find(path)
        except LookupError:
            nltk.download(name, quiet=True)


ensure_nltk_data()

STOP_WORDS = set(stopwords.words("english"))
LEMMATIZER = WordNetLemmatizer()
NEGATION_CUES = {"no", "not", "nor", "never", "without", "none"}


def clean_text(text: str) -> str:
    """Basic clean + lemmatize pipeline (mirrors the notebook's df2 cleaning)."""
    if not isinstance(text, str):
        return ""
    text = text.lower()
    text = re.sub(r"http\S+", "", text)
    text = re.sub(r"[^a-zA-Z ]", "", text)

    words = word_tokenize(text)
    words = [w for w in words if w not in STOP_WORDS]
    words = [LEMMATIZER.lemmatize(w) for w in words]
    return " ".join(words)


def clean_text_negation(text: str) -> str:
    """Clean + lemmatize while tagging words that follow a negation cue.

    A negated word gets prefixed with 'NEG_' and is excluded from
    plain keyword overlap, so "no SQL experience" won't be treated as a
    match for a job description that mentions SQL.
    """
    if not isinstance(text, str):
        return ""
    text = text.lower()
    sentences = re.split(r"[.!?]", text)
    processed_tokens = []

    for sent in sentences:
        words = re.findall(r"[a-z]+", sent)
        negate = False
        for w in words:
            if w in NEGATION_CUES:
                negate = True
                continue
            tagged = f"NEG_{w}" if negate else w
            processed_tokens.append(tagged)

    tokens = [
        ("NEG_" if t.startswith("NEG_") else "")
        + LEMMATIZER.lemmatize(t.replace("NEG_", ""))
        for t in processed_tokens
        if t.replace("NEG_", "") not in STOP_WORDS
        and len(t.replace("NEG_", "")) > 1
    ]
    return " ".join(tokens)


def rank_resumes(job_description: str, resumes: dict) -> "tuple":
    """Rank resumes against a job description using TF-IDF cosine similarity.

    Args:
        job_description: raw JD text.
        resumes: dict of {resume_name: raw_resume_text}.

    Returns:
        (results_df, clean_jd, clean_resumes) where results_df has columns
        ['resume', 'match_percent'], sorted descending.
    """
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import pandas as pd

    clean_jd = clean_text_negation(job_description)
    clean_resumes = {name: clean_text_negation(text) for name, text in resumes.items()}

    documents = [clean_jd] + list(clean_resumes.values())
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(documents)

    jd_vector = tfidf_matrix[0:1]
    resume_vectors = tfidf_matrix[1:]
    scores = cosine_similarity(jd_vector, resume_vectors)[0]

    results = pd.DataFrame(
        {
            "resume": list(clean_resumes.keys()),
            "match_percent": (scores * 100).round(1),
        }
    ).sort_values("match_percent", ascending=False).reset_index(drop=True)

    return results, clean_jd, clean_resumes


def matched_keywords(clean_jd: str, clean_resume: str, top_n: int = 10) -> list:
    """Return keywords shared between a cleaned JD and a cleaned resume.

    Negated terms (NEG_*) are excluded on both sides so a rejected/negated
    mention doesn't get shown as a "matching" skill.
    """
    jd_terms = {t for t in clean_jd.split() if not t.startswith("NEG_")}
    resume_terms = {t for t in clean_resume.split() if not t.startswith("NEG_")}
    return sorted(jd_terms & resume_terms)[:top_n]


# ---------------------------------------------------------------------------
# File text extraction: lets the app accept .txt, .pdf, .docx, and image
# (.jpg/.jpeg/.png/etc.) resume uploads, not just pasted text.
# ---------------------------------------------------------------------------


def _extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF. Falls back to OCR for scanned/image-only PDFs."""
    text_parts = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    except Exception:
        pass

    text = "\n".join(text_parts).strip()
    if text:
        return text

    # Likely a scanned/image-based PDF with no embedded text layer — try OCR.
    try:
        from pdf2image import convert_from_bytes

        images = convert_from_bytes(file_bytes)
        ocr_parts = [_extract_text_from_image_obj(img) for img in images]
        return "\n".join(ocr_parts).strip()
    except Exception:
        return ""


def _extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs]

    # Also pull text out of any tables in the document.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def _extract_text_from_image_obj(image) -> str:
    import pytesseract

    return pytesseract.image_to_string(image)


def _extract_text_from_image_bytes(file_bytes: bytes) -> str:
    from PIL import Image

    image = Image.open(io.BytesIO(file_bytes))
    return _extract_text_from_image_obj(image)


def extract_text_from_file(uploaded_file) -> str:
    """Extract raw text from an uploaded file of any supported type.

    Args:
        uploaded_file: a Streamlit UploadedFile (has .name and .read()/.getvalue()).

    Returns:
        Extracted text, or "" if extraction failed or the type is unsupported.

    Raises:
        ValueError: if the file extension isn't one we support, so the caller
        can show a clear message instead of silently returning nothing.
    """
    name = getattr(uploaded_file, "name", "file")
    ext = os.path.splitext(name)[1].lower()

    file_bytes = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()

    if ext == ".txt":
        return _extract_text_from_txt(file_bytes)
    elif ext == ".pdf":
        return _extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return _extract_text_from_docx(file_bytes)
    elif ext in IMAGE_EXTENSIONS:
        return _extract_text_from_image_bytes(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}' for {name}. "
            "Supported: .txt, .pdf, .docx, .jpg, .jpeg, .png"
        )

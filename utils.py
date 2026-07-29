"""
Shared text-processing + matching utilities for the Resume Screening System.

Two cleaning paths are kept, matching the two things the model was trained
around in the original notebook:

1. clean_text()            -> simple clean/lemmatize, used to prep text
                               before TF-IDF for the SVM classifier.
2. clean_text_negation()   -> same idea but tags words that follow a
                               negation cue (e.g. "not fluent in Java" ->
                               "NEG_fluent NEG_in NEG_java"), used for the
                               JD-vs-resume matching/ranking feature so that
                               "no experience with X" doesn't score as a
                               match for X.
"""
import io
import os
import re

import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"}

def ensure_nltk_data() -> None:
    """Download required NLTK corpora if they aren't already present."""
    resources = {
        "corpora/stopwords": "stopwords",
        "corpora/wordnet": "wordnet",
        "tokenizers/punkt": "punkt",
        "tokenizers/punkt_tab": "punkt_tab",
    }
    for path, name in resources.items():
        try:
            nltk.data.find(path)
        except LookupError:
            nltk.download(name, quiet=True)

ensure_nltk_data()

STOP_WORDS = set(stopwords.words("english"))
LEMMATIZER = WordNetLemmatizer()
NEGATION_CUES = {"no", "not", "nor", "never", "without", "none"}

def clean_text(text: str) -> str:
    """Basic clean + lemmatize pipeline (mirrors the notebook's df2 cleaning)."""
    if not isinstance(text, str):
        return ""
    text = text.lower()
    text = re.sub(r"http\S+", "", text)
    text = re.sub(r"[^a-zA-Z ]", "", text)

    words = word_tokenize(text)
    words = [w for w in words if w not in STOP_WORDS]
    words = [LEMMATIZER.lemmatize(w) for w in words]
    return " ".join(words)

def clean_text_negation(text: str) -> str:
    """Clean + lemmatize while tagging words that follow a negation cue.

    A negated word gets prefixed with 'NEG_' and is excluded from
    plain keyword overlap, so "no SQL experience" won't be treated as a
    match for a job description that mentions SQL.
    """
    if not isinstance(text, str):
        return ""
    text = text.lower()
    sentences = re.split(r"[.!?]", text)
    processed_tokens = []

    for sent in sentences:
        words = re.findall(r"[a-z]+", sent)
        negate = False
        for w in words:
            if w in NEGATION_CUES:
                negate = True
                continue
            tagged = f"NEG_{w}" if negate else w
            processed_tokens.append(tagged)

    tokens = [
        ("NEG_" if t.startswith("NEG_") else "")
        + LEMMATIZER.lemmatize(t.replace("NEG_", ""))
        for t in processed_tokens
        if t.replace("NEG_", "") not in STOP_WORDS
        and len(t.replace("NEG_", "")) > 1
    ]
    return " ".join(tokens)

def rank_resumes(job_description: str, resumes: dict) -> "tuple":
    """Rank resumes against a job description using TF-IDF cosine similarity.

    Args:
        job_description: raw JD text.
        resumes: dict of {resume_name: raw_resume_text}.

    Returns:
        (results_df, clean_jd, clean_resumes) where results_df has columns
        ['resume', 'match_percent'], sorted descending.
    """
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import pandas as pd

    clean_jd = clean_text_negation(job_description)
    clean_resumes = {name: clean_text_negation(text) for name, text in resumes.items()}

    documents = [clean_jd] + list(clean_resumes.values())
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(documents)

    jd_vector = tfidf_matrix[0:1]
    resume_vectors = tfidf_matrix[1:]
    scores = cosine_similarity(jd_vector, resume_vectors)[0]

    results = pd.DataFrame(
        {
            "resume": list(clean_resumes.keys()),
            "match_percent": (scores * 100).round(1),
        }
    ).sort_values("match_percent", ascending=False).reset_index(drop=True)

    return results, clean_jd, clean_resumes

def matched_keywords(clean_jd: str, clean_resume: str, top_n: int = 10) -> list:
    """Return keywords shared between a cleaned JD and a cleaned resume.

    Negated terms (NEG_*) are excluded on both sides so a rejected/negated
    mention doesn't get shown as a "matching" skill.
    """
    jd_terms = {t for t in clean_jd.split() if not t.startswith("NEG_")}
    resume_terms = {t for t in clean_resume.split() if not t.startswith("NEG_")}
    return sorted(jd_terms & resume_terms)[:top_n]

# --
# File text extraction: lets the app accept .txt, .pdf, .docx, and image
# (.jpg/.jpeg/.png/etc.) resume uploads, not just pasted text.
# --
def _extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")

def _extract_text_from_pdf_with_meta(file_bytes: bytes) -> "tuple":
    """Extract text from a PDF, plus whether OCR had to be used.

    Falls back to OCR for scanned/image-only PDFs. Raises RuntimeError with a
    specific reason if both the direct extraction and the OCR fallback fail
    to produce any text, so the caller can show the real cause instead of a
    generic "no text found" message.

    Returns:
        (text, used_ocr)
    """
    text_parts = []
    text_error = None
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    except Exception as e:
        text_error = e

    text = "\n".join(text_parts).strip()
    if text:
        return text, False

    # No embedded text layer (or reading failed) — likely a scanned/image
    # PDF. Try OCR via pdf2image (needs the Poppler system package) +
    # pytesseract (needs the Tesseract system package).
    ocr_error = None
    try:
        from pdf2image import convert_from_bytes

        images = convert_from_bytes(file_bytes)
        ocr_text = "\n".join(_extract_text_from_image_obj(img) for img in images).strip()
        if ocr_text:
            return ocr_text, True
    except Exception as e:
        ocr_error = e

    # Both paths failed to produce text — surface the real reason.
    if ocr_error is not None:
        raise RuntimeError(
            "This looks like a scanned/image PDF with no text layer, and OCR "
            f"failed: {ocr_error}. This usually means the 'poppler' and/or "
            "'tesseract' system packages aren't installed (see README's "
            "'Run locally' / packages.txt sections)."
        )
    if text_error is not None:
        raise RuntimeError(f"Couldn't read this PDF: {text_error}")
    raise RuntimeError(
        "This PDF appears to have no extractable text (empty or scanned "
        "pages with no OCR result)."
    )

def _extract_text_from_docx_with_meta(file_bytes: bytes) -> "tuple":
    """Extract text from a DOCX, plus whether it contains tables.

    Returns:
        (text, has_tables)
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs]

    has_tables = len(doc.tables) > 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs), has_tables


def _extract_text_from_image_obj(image) -> str:
    import pytesseract

    return pytesseract.image_to_string(image)

def _extract_text_from_image_bytes(file_bytes: bytes) -> str:
    from PIL import Image

    image = Image.open(io.BytesIO(file_bytes))
    return _extract_text_from_image_obj(image)

def extract_text_and_meta(uploaded_file) -> "tuple":
    """Extract text plus format metadata used by the ATS Score Checker.

    Returns:
        (text, meta) where meta is a dict with keys:
            - ext: file extension, e.g. '.pdf'
            - used_ocr: True if OCR had to be used (PDF had no text layer,
              or the file was an image) — real-world ATS parsers generally
              can't OCR, so this is a strong red flag for parseability.
            - has_tables: True if a DOCX file contains tables (many ATS
              parsers scramble or skip table content).

    Raises:
        ValueError: unsupported file extension.
        RuntimeError: extraction failed (see message for the reason).
    """
    name = getattr(uploaded_file, "name", "file")
    ext = os.path.splitext(name)[1].lower()
    file_bytes = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()

    meta = {"ext": ext, "used_ocr": False, "has_tables": False}

    if ext == ".txt":
        return _extract_text_from_txt(file_bytes), meta
    elif ext == ".pdf":
        text, used_ocr = _extract_text_from_pdf_with_meta(file_bytes)
        meta["used_ocr"] = used_ocr
        return text, meta
    elif ext == ".docx":
        text, has_tables = _extract_text_from_docx_with_meta(file_bytes)
        meta["has_tables"] = has_tables
        return text, meta
    elif ext in IMAGE_EXTENSIONS:
        meta["used_ocr"] = True
        return _extract_text_from_image_bytes(file_bytes), meta
    else:
        raise ValueError(
            f"Unsupported file type '{ext}' for {name}. "
            "Supported: .txt, .pdf, .docx, .jpg, .jpeg, .png"
        )

def extract_text_from_file(uploaded_file) -> str:
    """Extract raw text from an uploaded file of any supported type.

    Args:
        uploaded_file: a Streamlit UploadedFile (has .name and .read()/.getvalue()).

    Returns:
        Extracted text.

    Raises:
        ValueError: if the file extension isn't one we support, so the caller
        can show a clear message instead of silently returning nothing.
        RuntimeError: if extraction genuinely failed (see message for why).
    """
    text, _meta = extract_text_and_meta(uploaded_file)
    return text


# --
# ATS Score Checker
# --
EMAIL_PATTERN = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_PATTERN = re.compile(
    r"(\+?\d{1,3}[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}"
)

CORE_SECTIONS = {
    "experience": ["experience", "work history", "employment"],
    "education": ["education", "academic"],
    "skills": ["skills", "technical skills", "core competencies"],
}
BONUS_SECTIONS = {
    "summary": ["summary", "objective", "profile"],
    "projects": ["projects"],
    "certifications": ["certifications", "certificates", "licenses"],
}

def extract_contact_info(text: str) -> dict:
    """Return {'email': str|None, 'phone': str|None} found in resume text."""
    email_match = EMAIL_PATTERN.search(text or "")
    phone_match = PHONE_PATTERN.search(text or "")
    return {
        "email": email_match.group(0) if email_match else None,
        "phone": phone_match.group(0) if phone_match else None,
    }

def check_sections(text: str) -> dict:
    """Return {section_name: bool_found} for core + bonus resume sections."""
    lower = (text or "").lower()
    found = {}
    for section, aliases in {**CORE_SECTIONS, **BONUS_SECTIONS}.items():
        found[section] = any(alias in lower for alias in aliases)
    return found

def compute_ats_score(resume_text: str, job_description: str, meta: dict = None) -> dict:
    """Compute an ATS-style compatibility score for a resume against a JD.

    Args:
        resume_text: raw resume text.
        job_description: raw job description text.
        meta: optional dict from extract_text_and_meta() with 'used_ocr' /
            'has_tables' flags, used to flag formatting/parseability issues.

    Returns:
        dict with keys:
            total (float, 0-100), rating (str), breakdown (dict of
            {label: (points_earned, points_possible)}), missing_keywords
            (list), suggestions (list), contact (dict), sections (dict),
            word_count (int).
    """
    meta = meta or {}
    breakdown = {}
    suggestions = []

    # 1. Keyword match with the job description (up to 50 pts)
    results, clean_jd, clean_resumes = rank_resumes(job_description, {"resume": resume_text})
    match_percent = float(results.loc[results["resume"] == "resume", "match_percent"].iloc[0])
    keyword_pts = round(match_percent / 100 * 50, 1)
    breakdown["Keyword match with job description"] = (keyword_pts, 50)

    jd_terms = {t for t in clean_jd.split() if not t.startswith("NEG_")}
    resume_terms = {t for t in clean_resumes["resume"].split() if not t.startswith("NEG_")}
    missing_keywords = sorted(jd_terms - resume_terms)[:15]
    if missing_keywords:
        suggestions.append(
            "Work these job-description keywords into your resume where "
            "genuinely true of your experience: " + ", ".join(missing_keywords[:10])
        )

    # 2. Contact information (up to 10 pts)
    contact = extract_contact_info(resume_text)
    contact_pts = (5 if contact["email"] else 0) + (5 if contact["phone"] else 0)
    breakdown["Contact information"] = (contact_pts, 10)
    if not contact["email"]:
        suggestions.append("Add an email address — many ATS systems use it as the primary candidate identifier.")
    if not contact["phone"]:
        suggestions.append("Add a phone number.")

    # 3. Standard resume sections (up to 20 pts, core sections only)
    sections = check_sections(resume_text)
    core_found = sum(sections[s] for s in CORE_SECTIONS)
    section_pts = round(core_found / len(CORE_SECTIONS) * 20, 1)
    breakdown["Standard resume sections"] = (section_pts, 20)
    missing_core = [s for s in CORE_SECTIONS if not sections[s]]
    if missing_core:
        suggestions.append(
            "Add clear section headers for: " + ", ".join(missing_core)
            + " — ATS parsers rely on standard headers to categorize your content."
        )

    # 4. ATS-friendly formatting (up to 10 pts)
    format_pts = 10.0
    format_notes = []
    if meta.get("used_ocr"):
        format_pts -= 7
        format_notes.append(
            "This looks like a scanned image rather than a text-based document. "
            "Most real ATS systems can only read embedded text, not images — "
            "they may see this resume as blank. Export from Word/Google Docs "
            "as a text-based PDF instead of scanning a printed copy."
        )
    if meta.get("has_tables"):
        format_pts -= 3
        format_notes.append(
            "Tables were detected. Many ATS parsers scramble or drop table "
            "content — use plain paragraphs and bullet lists instead."
        )
    format_pts = max(format_pts, 0)
    breakdown["ATS-friendly formatting"] = (round(format_pts, 1), 10)
    suggestions.extend(format_notes)

    # 5. Resume length (up to 10 pts) — rough heuristic, ~350-900 words is typical for 1-2 pages
    word_count = len((resume_text or "").split())
    if 300 <= word_count <= 900:
        length_pts = 10.0
    elif word_count < 300:
        length_pts = round(max(word_count, 0) / 300 * 10, 1)
        suggestions.append(
            f"Resume is quite short ({word_count} words) — ATS and recruiters "
            "typically expect more detail on experience and skills."
        )
    else:
        length_pts = max(10 - (word_count - 900) / 100, 0)
        suggestions.append(
            f"Resume is quite long ({word_count} words) — consider trimming to "
            "the most relevant experience for this role."
        )
    breakdown["Resume length"] = (round(length_pts, 1), 10)

    total = round(sum(pts for pts, _ in breakdown.values()), 1)
    if total >= 80:
        rating = "Strong match"
    elif total >= 60:
        rating = "Moderate match"
    elif total >= 40:
        rating = "Weak match"
    else:
        rating = "Poor match"

    return {
        "total": total,
        "rating": rating,
        "breakdown": breakdown,
        "missing_keywords": missing_keywords,
        "suggestions": suggestions,
        "contact": contact,
        "sections": sections,
        "word_count": word_count,
    }
